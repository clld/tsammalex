from itertools import chain, groupby
import os
from datetime import date

from sqlalchemy.orm import joinedload, contains_eager
from six import BytesIO
from six.moves.urllib.request import urlopen
from docx import Document
from docx.shared import Inches
try:
    from xhtml2pdf import pisa
except ImportError:
    class pisa(object):
        @staticmethod
        def CreatePDF(*args, **kw):
            print("ERROR: xhtml2pdf is not installed!")
from path import path

from clld.web.util.helpers import text_citation, charis_font_spec_css
from clld.web.adapters import get_adapter
from clld.web.adapters.geojson import (
    GeoJsonParameterMultipleValueSets, GeoJson, GeoJsonLanguages,
)
from clld.web.adapters.base import Representation
from clld.web.adapters.download import Download
from clld.interfaces import ILanguage, IIndex, IRepresentation, IParameter
from clld.db.meta import DBSession
from clld.db.models.common import Parameter, Language, ValueSet

import tsammalex
from tsammalex.interfaces import IEcoregion
from tsammalex.models import Taxon


download_path = lambda basename: \
    path(tsammalex.__file__).dirname().joinpath('static', 'download', basename)


css_tmpl = """
    {0}

    html,body {{
        font-family: 'charissil'; font-size: 3.5mm;
    }}
    @page title_template {{ margin-top: 5cm; }}
    @page regular_template {{
        size: a4 portrait;
        @frame header_frame {{           /* Static Frame */
            -pdf-frame-content: header_content;
            left: 40pt; width: 532pt; top: 40pt; height: 30pt;
        }}
        @frame content_frame {{          /* Content Frame */
            left: 40pt; width: 532pt; top: 80pt; height: 652pt;
        }}
        @frame footer_frame {{           /* Another static Frame */
            -pdf-frame-content: footer_content;
            left: 40pt; width: 532pt; top: 782pt; height: 20pt;
        }}
    }}
    div.title {{ margin-bottom: 5cm; }}
    h1 {{ font-size: 30mm; text-align: center; }}
    h2 {{ -pdf-keep-with-next: true; padding-bottom: -2mm; }}
    h3 {{ -pdf-keep-with-next: true; padding-bottom: -2mm; }}
    p {{ -pdf-keep-with-next: true; }}
    p.separator {{ -pdf-keep-with-next: false; font-size: 1mm; }}
    img.image {{ zoom: 55%; }}
    td {{ text-align: center; }}
"""

html_tmpl = """
<html><head><style>%s</style></head><body>
    <div id="header_content" style="text-align: center;">%s</div>

    <div id="footer_content" style="text-align: center;">
        <pdf:pagenumber> of <pdf:pagecount>
    </div>
    <pdf:nexttemplate name="title_template" />
    <p>&nbsp;<p>
    <p>&nbsp;<p>
    <p>&nbsp;<p>
    <p>&nbsp;<p>
    <div class="title">
    %s
    </div>
    <pdf:nexttemplate name="regular_template" />
    <pdf:nextpage />
    %s
    </body></html>
"""


class Pdf(Download):
    ext = 'pdf'
    description = "printable PDF file"

    def asset_spec(self, req):
        return '.'.join(Download.asset_spec(self, req).split('.')[:-1])

    def create(self, req, filename=None, verbose=True, link_callback=None, lang=None):
        html = []
        lang = lang or Language.get('afr')
        entries = list(
            DBSession.query(ValueSet).join(ValueSet.parameter)
            .filter(ValueSet.language_pk == lang.pk)
            .order_by(
                Taxon.kingdom,
                Taxon.phylum,
                Taxon.class_,
                Taxon.order,
                Taxon.family,
                Parameter.name)
            .options(contains_eager(ValueSet.parameter), joinedload(ValueSet.values)))

        for kingdom, taxa1 in groupby(entries, key=lambda vs: vs.parameter.kingdom):
            html.append('<h2>Kingdom: %s</h2>' % (kingdom or 'other'))
            for phylum, taxa2 in groupby(taxa1, key=lambda vs: vs.parameter.phylum):
                html.append('<h3>Phylum: %s</h3>' % (phylum or 'other'))
                for class_, taxa3 in groupby(taxa2, key=lambda vs: vs.parameter.class_):
                    html.append('<h4>Class: %s</h4>' % (class_ or 'other'))
                    for order, taxa4 in groupby(taxa3, key=lambda vs: vs.parameter.order):
                        html.append('<h5>Order: %s</h5>' % (order or 'other'))
                        for family, taxa5 in groupby(taxa4, key=lambda vs: vs.parameter.family):
                            html.append('<h6>Family: %s</h6>' % (family or 'other'))
                            for entry in taxa5:
                                adapter = get_adapter(
                                    IRepresentation, entry, req, ext='snippet.html')
                                html.append(adapter.render(entry, req))
                                html.append('<p class="separator">&nbsp;<p>')

        with open(download_path('%s.pdf' % lang.id), 'wb') as fp:
            editors = ''
            if lang.contribution.contributor_assocs:
                editors = 'edited by ' + ' and '.join(
                    c.last_first() for c in lang.contribution.primary_contributors)
            pisa.CreatePDF(
                html_tmpl % (
                    css_tmpl.format(charis_font_spec_css()),
                    req.resource_url(req.dataset),
                    """
<h1 style="text-align: center; font-size: 12mm;">%(language)s names for Plants and Animals</h1>
<h2 style="text-align: center; font-size: 8mm;">%(editors)s</h2>
<p style="font-size: 5mm;">
This document was created from <a href="%(url)s">%(dataset)s</a> on %(date)s.
</p>
<p style="font-size: 5mm;">
%(dataset)s is published under a %(license)s and should be cited as
</p>
<blockquote style="font-size: 5mm;"><i>%(citation)s</i></blockquote>
<p style="font-size: 5mm;">
A full list of contributors is available at
<a href="%(url)scontributors">%(url)scontributors</a>
</p>
<p style="font-size: 5mm;">
The list of references cited in this document is available at
<a href="%(url)ssources">%(url)ssources</a>
</p>
""" % dict(
                    language=lang.name,
                    editors=editors,
                    dataset=req.dataset.name,
                    url=req.resource_url(req.dataset),
                    date=date.today(),
                    citation=text_citation(req, req.dataset),
                    license=req.dataset.jsondata['license_name']),
                    ''.join(html)),
                dest=fp,
                link_callback=link_callback,
            )


class GeoJsonEcoregions(GeoJson):
    def featurecollection_properties(self, ctx, req):
        return {'name': "WWF's Terrestrial Ecoregions of the Afrotropics"}

    def get_features(self, ctx, req):
        for ecoregion in ctx.get_query():
            for polygon in ecoregion.jsondata['polygons']:
                yield {
                    'type': 'Feature',
                    'properties': {
                        'id': ecoregion.id,
                        'label': '%s %s' % (ecoregion.id, ecoregion.name),
                        'color': ecoregion.biome.description,
                        'language': {'id': ecoregion.id},
                        'latlng': [ecoregion.latitude, ecoregion.longitude],
                    },
                    'geometry': polygon,
                }


class GeoJsonTaxa(GeoJsonParameterMultipleValueSets):
    def feature_properties(self, ctx, req, p):
        return {
            'lineage': p[0].lineage,
            'label': ', '.join(v.name for v in chain(*[vs.values for vs in p[1]]))}


class GeoJsonLanguoids(GeoJsonLanguages):
    def feature_properties(self, ctx, req, feature):
        return {'lineage': feature.lineage}


class LanguagePdf(Representation):
    mimetype = 'application/pdf'
    extension = 'pdf'

    def render(self, ctx, req):
        fname = download_path('%s.pdf' % ctx.id)
        if fname.exists():
            return fname.bytes()


class Docx(Representation):
    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    extension = 'docx'

    def write(self, ctx, req, doc):
        raise NotImplemented  # pragma: no cover

    def render(self, ctx, req):
        document = Document()
        self.write(ctx, req, document)
        #
        # TODO: add tsammalex license information!
        #
        d = BytesIO()
        document.save(d)
        d.seek(0)
        return d.read()


class LanguageDocx(Docx):
    def write(self, ctx, req, document):
        document.add_heading(ctx.name, 0)
        table = document.add_table(rows=len(ctx.valuesets) + 1, cols=2)
        table.cell(0, 0).text = 'Taxon'
        table.cell(0, 1).text = 'Name'

        for i, vs in enumerate(ctx.valuesets):
            table.cell(i + 1, 0).text = vs.parameter.name
            table.cell(i + 1, 1).text = ', '.join(v.name for v in vs.values)


class TaxonDocx(Docx):
    def write(self, ctx, req, document):
        document.add_heading(ctx.name, 0)
        document.add_heading('Names', 1)
        table = document.add_table(rows=len(ctx.valuesets), cols=2)

        for i, vs in enumerate(ctx.valuesets):
            table.cell(i, 0).text = vs.language.name
            table.cell(i, 1).text = ', '.join(v.name for v in vs.values)

        document.add_heading('Photos', 1)
        for f in ctx._files:
            try:
                stream = BytesIO(urlopen(f.jsondata['url']).read())
            except:  # pragma: no cover
                continue
            document.add_picture(stream, width=Inches(3.5))

            table = document.add_table(rows=0, cols=2)
            for attr in 'date place creator source permission comments'.split():
                v = f.get_data(attr)
                if v:
                    cells = table.add_row().cells
                    cells[0].text = attr.capitalize()
                    cells[1].text = v

            # p = document.add_paragraph('A plain paragraph having some ')
            # p.add_run('bold').bold = True
            # p.add_run(' and some ')
            # p.add_run('italic.').italic = True
            #
            # document.add_heading('Heading, level 1', level=1)
            # document.add_paragraph('Intense quote', style='IntenseQuote')
            #
            # document.add_paragraph(
            #     'first item in unordered list', style='ListBullet'
            # )
            # document.add_paragraph(
            #     'first item in ordered list', style='ListNumber'
            # )
            #
            # #document.add_picture('monty-truth.png', width=Inches(1.25))
            #
            # table = document.add_table(rows=1, cols=3)
            # hdr_cells = table.rows[0].cells
            # hdr_cells[0].text = 'Qty'
            # hdr_cells[1].text = 'Id'
            # hdr_cells[2].text = 'Desc'
            # #for item in recordset:
            # #    row_cells = table.add_row().cells
            # #    row_cells[0].text = str(item.qty)
            # #    row_cells[1].text = str(item.id)
            # #    row_cells[2].text = item.desc
            #
            # document.add_page_break()


def includeme(config):
    config.register_adapter(LanguagePdf, ILanguage)
    #config.register_adapter(LanguageDocx, ILanguage)
    config.register_adapter(TaxonDocx, IParameter)
    config.register_adapter(GeoJsonEcoregions, IEcoregion, IIndex)
    config.register_adapter(GeoJsonTaxa, IParameter)
    config.register_adapter(GeoJsonLanguoids, ILanguage, IIndex)
